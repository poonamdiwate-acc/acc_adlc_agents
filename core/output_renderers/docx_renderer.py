"""DOCX output renderer — generates Word documents for all agents."""

from __future__ import annotations

import io
import logging
from typing import Any, Dict, List

from core.output_renderers.base import OutputRenderer, RenderedOutput

logger = logging.getLogger(__name__)


class DocxRenderer(OutputRenderer):
    """Renders result as a structured Word document."""

    @property
    def format_name(self) -> str:
        return "docx"

    def render(
        self,
        result: Dict[str, Any],
        *,
        agent_id: str,
        run_id: str,
    ) -> RenderedOutput:
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError as exc:
            raise RuntimeError(
                "python-docx is required for DOCX rendering. "
                "Install with: pip install python-docx"
            ) from exc

        doc = Document()

        if "nfr_specifications" in result or "security_controls" in result:
            self._render_de06(doc, result, agent_id, run_id)
        elif "cost_estimate" in result or "optimization_plan" in result:
            self._render_de08(doc, result, agent_id, run_id)
        elif "gap_report" in result or "gap_summary" in result:
            self._render_ad04(doc, result, agent_id, run_id)
        else:
            self._render_generic(doc, result, agent_id, run_id)

        buffer = io.BytesIO()
        doc.save(buffer)
        content = buffer.getvalue()

        return RenderedOutput(
            content=content,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{agent_id}_output_{run_id}.docx",
        )

    # ------------------------------------------------------------------
    # DE-06: Non-Functional Design
    # ------------------------------------------------------------------

    def _render_de06(
        self, doc: Any, result: Dict[str, Any], agent_id: str, run_id: str
    ) -> None:
        doc.add_heading(f"{agent_id} — Non-Functional Design Report", level=0)
        doc.add_paragraph(f"Run ID: {run_id}")

        nfr_specs: List[Dict[str, Any]] = result.get("nfr_specifications", [])
        if nfr_specs:
            doc.add_heading("NFR Specifications", level=1)
            doc.add_paragraph(f"Total NFRs identified: {len(nfr_specs)}")

            table = doc.add_table(rows=1, cols=7)
            table.style = "Table Grid"
            headers = ["NFR ID", "Name", "Category", "Target / Threshold", "Priority", "Confidence", "Rationale"]
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h

            for nfr in nfr_specs:
                row = table.add_row().cells
                row[0].text = nfr.get("nfr_id", "")
                row[1].text = nfr.get("nfr_name", nfr.get("title", ""))
                row[2].text = nfr.get("category", "")
                row[3].text = f"{nfr.get('target_metric', '')} — {nfr.get('threshold', '')}"
                row[4].text = nfr.get("priority", "")
                row[5].text = nfr.get("confidence", "")
                row[6].text = nfr.get("rationale", "")

        security = result.get("security_controls", {})
        if security:
            doc.add_heading("Security Controls", level=1)

            posture = security.get("overall_posture", "N/A")
            recommendation = security.get("recommendation", "N/A")
            doc.add_paragraph(f"Overall Posture: {posture} | Recommendation: {recommendation}")

            threat_summary = security.get("threat_surface_summary", "")
            if threat_summary:
                doc.add_heading("Threat Surface Summary", level=2)
                doc.add_paragraph(threat_summary)

            controls: List[Dict[str, Any]] = security.get("controls", [])
            if controls:
                doc.add_heading("Controls Matrix", level=2)
                table = doc.add_table(rows=1, cols=6)
                table.style = "Table Grid"
                headers = ["Control ID", "Domain", "Name", "Mechanism", "Confidence", "Priority"]
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h

                for ctrl in controls:
                    row = table.add_row().cells
                    row[0].text = ctrl.get("control_id", "")
                    row[1].text = ctrl.get("domain", "")
                    row[2].text = ctrl.get("name", "")
                    row[3].text = ctrl.get("mechanism", "")
                    row[4].text = ctrl.get("confidence", "")
                    row[5].text = ctrl.get("priority", "")

            mappings: List[Dict[str, Any]] = security.get("compliance_mappings", [])
            if mappings:
                doc.add_heading("Compliance Mappings", level=2)
                for m in mappings:
                    doc.add_paragraph(
                        f"{m.get('standard', 'N/A')}: {', '.join(m.get('applicable_controls', []))}"
                    )

    # ------------------------------------------------------------------
    # DE-08: Cost & Optimization
    # ------------------------------------------------------------------

    def _render_de08(
        self, doc: Any, result: Dict[str, Any], agent_id: str, run_id: str
    ) -> None:
        doc.add_heading(f"{agent_id} — Cost & Optimization Report", level=0)
        doc.add_paragraph(f"Run ID: {run_id}")

        cost_est = result.get("cost_estimate", {})
        if cost_est:
            doc.add_heading("Cost Estimate Summary", level=1)
            total_monthly = cost_est.get("total_monthly_usd", 0)
            total_annual = cost_est.get("total_annual_usd", 0)
            doc.add_paragraph(
                f"Total Monthly: ${total_monthly:,.2f} | "
                f"Total Annual: ${total_annual:,.2f} | "
                f"Confidence: {cost_est.get('overall_confidence', 'N/A')} | "
                f"Recommendation: {cost_est.get('recommendation', 'N/A')}"
            )

            assumptions = cost_est.get("assumptions", [])
            if assumptions:
                doc.add_heading("Assumptions", level=2)
                for a in assumptions:
                    doc.add_paragraph(f"• {a}")

            line_items: List[Dict[str, Any]] = cost_est.get("line_items", [])
            if line_items:
                doc.add_heading("Cost Line Items", level=2)
                table = doc.add_table(rows=1, cols=6)
                table.style = "Table Grid"
                headers = ["Cost ID", "Service", "Category", "Monthly USD", "Confidence", "Rationale"]
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h

                for item in line_items:
                    row = table.add_row().cells
                    row[0].text = item.get("cost_id", "")
                    row[1].text = item.get("service", "")
                    row[2].text = item.get("category", "")
                    row[3].text = f"${item.get('monthly_usd', 0):,.2f}"
                    row[4].text = item.get("confidence", "")
                    row[5].text = item.get("rationale", "")

        opt_plan: List[Dict[str, Any]] = result.get("optimization_plan", [])
        if opt_plan:
            doc.add_heading("Optimization Plan", level=1)
            table = doc.add_table(rows=1, cols=7)
            table.style = "Table Grid"
            headers = ["Opt ID", "Title", "Category", "Savings %", "Savings $/mo", "Priority", "Trade-off"]
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h

            for opt in opt_plan:
                row = table.add_row().cells
                row[0].text = opt.get("opt_id", "")
                row[1].text = opt.get("title", "")
                row[2].text = opt.get("category", "")
                row[3].text = f"{opt.get('estimated_savings_pct', 0):.1f}%"
                row[4].text = f"${opt.get('estimated_savings_monthly_usd', 0):,.2f}"
                row[5].text = opt.get("priority", "")
                row[6].text = opt.get("trade_off", "")

    # ------------------------------------------------------------------
    # AD-04: Gap Detection
    # ------------------------------------------------------------------

    def _render_ad04(
        self, doc: Any, result: Dict[str, Any], agent_id: str, run_id: str
    ) -> None:
        doc.add_heading(f"{agent_id} — Gap Detection Report", level=0)
        doc.add_paragraph(f"Run ID: {run_id}")

        gap_summary = result.get("gap_summary", {})
        if gap_summary:
            doc.add_heading("Gap Summary", level=1)
            summary_table = doc.add_table(rows=1, cols=2)
            summary_table.style = "Table Grid"
            hdr_cells = summary_table.rows[0].cells
            hdr_cells[0].text = "Metric"
            hdr_cells[1].text = "Value"

            summary_fields = [
                ("Total Requirements Analysed", gap_summary.get("total_requirements_analysed", 0)),
                ("Total Gaps Found", gap_summary.get("total_gaps_found", 0)),
                ("Blocking Gaps", gap_summary.get("blocking_gaps", 0)),
                ("Overall Quality", gap_summary.get("overall_quality", "N/A")),
                ("Recommendation", gap_summary.get("recommendation", "N/A")),
            ]
            for label, value in summary_fields:
                row = summary_table.add_row().cells
                row[0].text = label
                row[1].text = str(value)

            by_severity = gap_summary.get("gaps_by_severity", {})
            if by_severity:
                doc.add_heading("Gaps by Severity", level=2)
                sev_table = doc.add_table(rows=1, cols=2)
                sev_table.style = "Table Grid"
                sev_table.rows[0].cells[0].text = "Severity"
                sev_table.rows[0].cells[1].text = "Count"
                for sev, count in by_severity.items():
                    row = sev_table.add_row().cells
                    row[0].text = sev
                    row[1].text = str(count)

        gap_report: List[Dict[str, Any]] = result.get("gap_report", [])
        if gap_report:
            doc.add_heading("Gap Report", level=1)
            table = doc.add_table(rows=1, cols=6)
            table.style = "Table Grid"
            headers = ["Gap ID", "Req Ref", "Type", "Severity", "Description", "Recommendation"]
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header

            for gap in gap_report:
                row = table.add_row().cells
                row[0].text = gap.get("gap_id", "")
                row[1].text = gap.get("req_id_ref") or "N/A"
                row[2].text = gap.get("gap_type", "")
                row[3].text = gap.get("severity", "")
                row[4].text = gap.get("description", "")
                row[5].text = gap.get("recommendation", "")
        else:
            doc.add_heading("Gap Report", level=1)
            doc.add_paragraph("No gaps detected — requirements are clean.")

    # ------------------------------------------------------------------
    # Generic fallback
    # ------------------------------------------------------------------

    def _render_generic(
        self, doc: Any, result: Dict[str, Any], agent_id: str, run_id: str
    ) -> None:
        import json

        doc.add_heading(f"{agent_id} — Agent Output", level=0)
        doc.add_paragraph(f"Run ID: {run_id}")
        doc.add_paragraph(json.dumps(result, indent=2, ensure_ascii=False))
