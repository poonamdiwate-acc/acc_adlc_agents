"""DOCX input parser — extracts text from free-form Word documents."""

from __future__ import annotations

import io
from typing import List

from core.input_parsers.base import InputParser, ParsedDocument


class DocxParser(InputParser):
    """Extracts text, headings, and tables from .docx files."""

    @property
    def supported_content_types(self) -> List[str]:
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ]

    def parse(self, content: bytes) -> ParsedDocument:
        if not content:
            raise ValueError("Empty DOCX input")

        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install python-docx"
            ) from exc

        try:
            doc = Document(io.BytesIO(content))
        except Exception as exc:
            raise ValueError(f"Invalid DOCX file: {exc}") from exc

        # Extract paragraphs grouped by headings
        sections: dict[str, str] = {}
        current_heading = "_preamble"
        current_text: list[str] = []
        full_text_parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            full_text_parts.append(text)

            if para.style and para.style.name.startswith("Heading"):
                # Save previous section
                if current_text:
                    sections[current_heading] = "\n".join(current_text)
                current_heading = text
                current_text = []
            else:
                current_text.append(text)

        # Save last section
        if current_text:
            sections[current_heading] = "\n".join(current_text)

        # Extract tables
        tables: list[list[list[str]]] = []
        for table in doc.tables:
            table_data: list[list[str]] = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append(table_data)

        raw_text = "\n\n".join(full_text_parts)
        if not raw_text.strip():
            raise ValueError("DOCX file contains no text content")

        return ParsedDocument(
            raw_text=raw_text,
            sections=sections,
            tables=tables,
            source_format="docx",
            metadata={
                "size_bytes": len(content),
                "paragraph_count": len(full_text_parts),
                "table_count": len(tables),
                "section_count": len(sections),
            },
        )
